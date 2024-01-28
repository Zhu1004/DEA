import docx
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.image.image import Image
from docx.parts.image import ImagePart
from docx.oxml.shape import CT_Picture
from PIL import Image
from io import BytesIO
from docx.shared import Pt
import os
import shutil

def copy_file(source_file, destination_path, new_filename):
    # 获取源文件的文件名
    filename = os.path.basename(source_file)
    # 构建目标文件的完整路径
    destination_file = os.path.join(destination_path, new_filename)

    shutil.copy2(source_file, destination_file)
def RemoveDir(filepath):#用于清空yolo——source文件
    if not os.path.exists(filepath):#如果不存在
        print("检查文件夹是否存在")
    else:#如果存在这样的文件夹就清空
        shutil.rmtree(filepath)
        os.mkdir(filepath)
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None
def get_picture(document: Document, paragraph: Paragraph):
    global count
    img = paragraph._element.xpath('.//pic:pic')
    if not img:
        return document
    img: CT_Picture = img[0]
    embed = img.xpath('.//a:blip/@r:embed')[0]
    # 检查图片是否为嵌入式图片
    if embed in document.part.related_parts:
        try:
            related_part: ImagePart = document.part.related_parts[embed]
            image_data = related_part._blob
            
            # 使用Pillow加载图像数据
            pil_image = Image.open(BytesIO(image_data))
            
            # 将图像转换为RGB模式
            pil_image = pil_image.convert('RGB')

            # 将图像转换为支持的格式（例如JPEG）
            if pil_image.format not in ['JPEG', 'PNG', 'GIF']:
                pil_image = pil_image.convert('RGB')
            path = './pdf/pic'

            # 检查保存路径是否存在，如果不存在则创建路径
            os.makedirs(os.path.dirname(path), exist_ok=True)
            # 创建一个文件来保存转换后的图像
            #RemoveDir('RedsealRomove/data/images')
            temp_image_path = f'RedsealRomove/data/images/{count}.jpg'
            pil_image.save(temp_image_path, format='JPEG')

            text = f"图片{count}"
            count+=1

            new_para = paragraph.insert_paragraph_before()
            # 在新段落里添加文本
            new_para.text = text
            # 删除本段
            delete_paragraph(paragraph)

        except Exception as e:
            print(f"处理图片时发生异常：{str(e)}")
    
    return document
def addpage(document):
    for paragraph in document.paragraphs:
        p = paragraph._p
        sectPrs = p.xpath("./w:pPr/w:sectPr")
        if sectPrs:
            for sectPr in sectPrs:
                num = sectPr.xpath('.//w:cols/@w:num')[0]
                if num == "1":
                    text = 'page'
                    # 添加新段落
                    new_para = paragraph.insert_paragraph_before(text)
                    # 在新段落里添加文本
                    #new_para.text = text
def run_1():
    global count
    #RemoveDir("wordimg/")
    RemoveDir('RedsealRomove/data/images')
    d = docx.Document("wordout/example.docx")
    count=1
    for i in d.paragraphs:
        i.paragraph_format.space_before = Pt(12)  # 设置段前间距为12磅
        d= get_picture(d,i)
    tar_dir = 'RedsealRomove/data/images'  # 指定目标文件夹

    if len(os.listdir(tar_dir)) == 0:  # 目标文件夹内容为空的情况下
        #print("转换完成！！！,结果位于wordout文件夹，而不是best.docx")
        print("无图片情况")
        source_file = 'wordout/example.docx'  # 源文件路径
        destination_path = ''  # 目标路径，不包括文件名
        new_filename = 'best.docx'  # 复制后的新文件名

        copy_file(source_file, destination_path, new_filename)
        d = docx.Document("best.docx")
        for i in d.paragraphs:
            i.paragraph_format.space_before = Pt(12)  # 设置段前间距为12磅
        addpage(d)
        d.save("best.docx")
        return 1

    #输出图片张数：
    print(f"提取图片{count}张")
    d.save("wordout/temp.docx")
    return 0
