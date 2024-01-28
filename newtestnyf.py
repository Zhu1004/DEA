from sentence_transformers import SentenceTransformer,util
import numpy as np
import re
from docx import Document
from sklearn.metrics.pairwise import cosine_similarity
from fuzzywuzzy import process
from tqdm import tqdm
import openpyxl
import difflib
import mysql.connector


def split_into_paragraphs(text):
    sentences = re.split(r'[~]+', text)
    return [s for s in sentences ]


def get_paragraph_page_number(paragraph, file_path,found_paragraphs):
    document = Document(file_path)
    paragraphs = document.paragraphs
    page_number = 1
    #found_paragraphs = set()

    for index, p in enumerate(paragraphs):
        if p.text == paragraph.text and index not in found_paragraphs:
            found_paragraphs.add(index)
            return page_number
        elif "page" == p.text:
            page_number += 1
    return -1


# 读取Word文档并按句切割
def read_and_tokenize_word_file(file_path,model):
    document = Document(file_path)
    paragraphs = document.paragraphs
    sentences_and_embeddings = []

    found_paragraphs = set()

    # 创建一个范围为段落数量的进度条
    progress_bar = tqdm(total=len(paragraphs), desc='正在进行文本向量嵌入', bar_format='{l_bar}{bar}')
    idx=0
    for paragraph in paragraphs:
        page_number = get_paragraph_page_number(paragraph, file_path, found_paragraphs)
        if page_number is None:
            page_number = -1

        text = paragraph.text
        sentences = split_into_paragraphs(text)
        #print(len(sentences))
        
        for sentence in sentences:

            embedding = model.encode([sentence])[0]
            sentences_and_embeddings.append((sentence, embedding, idx, page_number))
            idx+=1

        # 更新进度条
        progress_bar.update(1)

    # 关闭进度条
    progress_bar.close()

    return sentences_and_embeddings






# 输入一个句子，找到与其最相似的三个句子及其所在段落
def find_top_k_similar_sentences(input_sentence,sentences_and_embeddings,model, k=3):
    input_embedding = model.encode([input_sentence])[0]

    similarities = cosine_similarity([input_embedding], [embedding for _, embedding, _, _ in sentences_and_embeddings])[0]

    #print("相似度：",similarities)
    top_k_indices = np.argsort(similarities)[-k:][::-1]
    #print(top_k_indices)
    top_k_sentences_and_paragraphs = [(sentences_and_embeddings[idx][0], sentences_and_embeddings[idx][2], sentences_and_embeddings[idx][3]) for idx in
                                      top_k_indices]

    return top_k_sentences_and_paragraphs

# while True:
#     input_sentence = input("请输入要找的内容：")
#     if input_sentence=='end':
#         break
#     word_file_path = "best.docx"
#     sentences_and_embeddings,sentences_and_embeddings_mh = read_and_tokenize_word_file(word_file_path)
#     results = process.extractOne(input_sentence,[s[0] for s in sentences_and_embeddings_mh])
#     top_3_similar_sentences_and_paragraphs = find_top_k_similar_sentences(input_sentence,sentences_and_embeddings)
#
#     print("输入句子：", input_sentence)
#     print("最相似的三个段落和页码：")
#     for sentence, paragraph_idx, page_number in top_3_similar_sentences_and_paragraphs:
#         print(f"句子：{sentence}\n段落：{paragraph_idx}\n页码：{page_number}\n")
#
#     if len(results) == 0:
#         print("process模块模糊搜索结果为NULL")
#     else:
#         print("process模块模糊搜索结果：")
#
#         print('内容：',results[0])
#         print("置信度：", results[1])
#         match_s=results[0]
#         for sentence, embedding, idx, page_number in sentences_and_embeddings_mh:
#             if sentence == match_s:
#                 print(f"段落：{idx}，页码：{page_number}")
#     print()
#
#     res = difflib.get_close_matches(input_sentence, [s[0] for s in sentences_and_embeddings_mh], 1, cutoff=0.2)
#     if len(res)==0:
#         print("difflib模块模糊搜索结果为NULL")
#     else:
#         print("difflib模块模糊搜索结果：")
#         print("句子：",res[0])
#         match_ss = res[0]
#         for sentence, embedding, idx, page_number in sentences_and_embeddings_mh:
#             if sentence == match_ss:
#                 print(f"段落：{idx}，页码：{page_number}")
#
#     emb0 = model1.encode(input_sentence)
#     emb1 = model1.encode(top_3_similar_sentences_and_paragraphs[0][0])
#     emb2 = model1.encode(top_3_similar_sentences_and_paragraphs[1][0])
#     emb3 = model1.encode(top_3_similar_sentences_and_paragraphs[2][0])
#     emb4 = model1.encode(match_s)
#     emb5 = model1.encode(match_ss)
#     cos_sim1 = util.cos_sim(emb0, emb1)
#     cos_sim2 = util.cos_sim(emb0, emb2)
#     cos_sim3 = util.cos_sim(emb0, emb3)
#     cos_sim4 = util.cos_sim(emb0, emb4)
#     cos_sim5 = util.cos_sim(emb0, emb5)
#     rrr=[]
#     rrr.append((match_ss,cos_sim5))
#     rrr.append((top_3_similar_sentences_and_paragraphs[0][0], cos_sim1))
#     rrr.append((top_3_similar_sentences_and_paragraphs[1][0], cos_sim2))
#     rrr.append((top_3_similar_sentences_and_paragraphs[2][0], cos_sim3))
#     rrr.append((match_s, cos_sim4))
#     sorted_list = sorted(rrr, key=lambda x: x[1], reverse=True)
#     print()
#     print("五个句子中与输入句子相似度最高的是：",rrr[0][0])
#     print("相似度：",rrr[0][1])

def run_5(relative_path,round):
    relative_path=relative_path.replace("\\", "/")
    # 加载预训练的句子嵌入模型
    print("模型加载")
    model = SentenceTransformer("model/outmodel/model")
    model1 = SentenceTransformer("model/moka-ai/m3e-base")
    try:
        db = mysql.connector.connect(
            host='127.0.0.1',
            port='3306',
            user='root',
            password='111111',
            database='dea'
        )
    except:
        print("连接数据库失败！")

    #那些每个pdf对应的指标tag由最开始传到excel中，打开数据库，遍历每行，找到数据库中对应的关键词和逻辑填入
    word_file_path = "best.docx"
    sentences_and_embeddings= read_and_tokenize_word_file(word_file_path,model)
    in_filename = "data_input.xlsx"
    wb = openpyxl.load_workbook(in_filename)
    sheet = wb.active
    col_tag = sheet['B']#输入的细则tag
    col_in=sheet['A']#从模板数据库中拿出关键字
    col_res1 = sheet['C']#定位结果与页码
    col_res11 = sheet['D']
    col_res2 = sheet['E']
    col_res22 = sheet['F']
    col_res3= sheet['G']
    col_res33= sheet['H']
    col_res4 = sheet['I']
    col_res44 = sheet['J']
    col_res5 = sheet['K']
    col_res55 = sheet['L']
    col_zl1=sheet['M']#输入glm的指令
    col_zl2 = sheet['N']
    col_zl3 = sheet['O']
    col_zl4 = sheet['P']
    col_zl5 = sheet['Q']
    col_score1=sheet['R']#得分
    col_score2 = sheet['S']
    col_score3 = sheet['T']
    col_score4 = sheet['U']
    col_score5 = sheet['V']
    col_list=[]
    col_list.append((col_res1,col_res11,col_zl1))
    col_list.append((col_res2,col_res22,col_zl2))
    col_list.append((col_res3, col_res33, col_zl3))
    col_list.append((col_res4, col_res44, col_zl4))
    col_list.append((col_res5, col_res55, col_zl5))
    print(len(col_tag),"len")
    for i in range(len(col_tag)):
        try:
            cursor = db.cursor()
            sql="select keywords,SCORING_LOGIC from deatest where tag="+str(col_tag[i].value)+";"
            print(sql)
            cursor.execute(sql)

            data = cursor.fetchone()
            l = []
            if data:
                for j in data:
                    l.append(j)
                    print(j)
                lj=l[1]
            else:
                print("tag对应出现错误！！！")
                break
            col_in[i].value = l[0]
            #拿到细则和逻辑
            #在此处要根据输入的细则去模板库中根据tag拿到关键词和逻辑
            input_sentence=l[0]
            visible_page = []
            results = process.extractOne(input_sentence, [s[0] for s in sentences_and_embeddings])
            top_3_similar_sentences_and_paragraphs = find_top_k_similar_sentences(input_sentence,sentences_and_embeddings,model)

            #print("输入句子：", input_sentence)
            #print("最相似的五个段落和页码：")
            for sentence, paragraph_idx, page_number in top_3_similar_sentences_and_paragraphs:
                visible_page.append((sentence,page_number,paragraph_idx))

            #print("process模块模糊搜索结果：")
            #print('内容：', results[0])
            #print("置信度：", results[1])
            if len(results)==0:
                match_s="NULL"
                visible_page.append((match_s,0,0))
            else:
                match_s = results[0]
                for sentence, embedding, idx, page_number in sentences_and_embeddings:
                    if sentence == match_s:
                        visible_page.append((sentence,page_number,idx))
                        break
            #print()

            #print("difflib模块模糊搜索结果：")
            res = difflib.get_close_matches(input_sentence, [s[0] for s in sentences_and_embeddings], 1, cutoff=0.2)
            #print("句子：", res[0])
            if len(res)==0:
                match_ss="NULL"
                visible_page.append((match_ss,0,0))
            else:
                match_ss = res[0]
                for sentence, embedding, idx, page_number in sentences_and_embeddings:
                    if sentence == match_ss:
                        visible_page.append((sentence,page_number,idx))
                        break

            emb0 = model1.encode(input_sentence)
            emb1 = model1.encode(top_3_similar_sentences_and_paragraphs[0][0])
            emb2 = model1.encode(top_3_similar_sentences_and_paragraphs[1][0])
            emb3 = model1.encode(top_3_similar_sentences_and_paragraphs[2][0])
            emb4 = model1.encode(match_s)
            emb5 = model1.encode(match_ss)
            cos_sim1 = util.cos_sim(emb0, emb1)
            cos_sim2 = util.cos_sim(emb0, emb2)
            cos_sim3 = util.cos_sim(emb0, emb3)
            cos_sim4 = util.cos_sim(emb0, emb4)
            cos_sim5 = util.cos_sim(emb0, emb5)
            rrr = []
            rrr.append((match_ss, cos_sim5))
            rrr.append((top_3_similar_sentences_and_paragraphs[0][0], cos_sim1))
            rrr.append((top_3_similar_sentences_and_paragraphs[1][0], cos_sim2))
            rrr.append((top_3_similar_sentences_and_paragraphs[2][0], cos_sim3))
            rrr.append((match_s, cos_sim4))
            rrr = sorted(rrr, key=lambda x: x[1], reverse=True)

            for rid,col in enumerate(col_list):
                col[0][i].value = rrr[rid][0]
                for sentence, page_number, idx in visible_page:
                    if sentence == rrr[rid][0]:
                        col[1][i].value = page_number
                        p_idx = idx
                        break
                if p_idx >= len(sentences_and_embeddings) - 1:
                    p_idx = p_idx - 1
                    p_next = p_idx
                else:
                    p_next = p_idx + 1

                if rrr[rid][0] == "NULL":
                    col[2][i].value = "NULL"
                else:
                    col[2][i].value = str(
                        "内容：" + sentences_and_embeddings[p_idx][0] + sentences_and_embeddings[p_next][0] +"。"+lj)
            cursor = db.cursor()
            try:
                sql = "insert into convert_result values("+str(col_tag[i].value)+",'"+relative_path+"',"+str(col_res11[i].value)+",'"+str(col_res1[i].value)+"');"
                print(sql)
                cursor.execute(sql)
            except:
                print("更新results操作")
                sql = "update convert_result set page="+str(col_res11[i].value)+" and results ='"+str(col_res1[i].value)+"' where tag=" + str(
                    col_tag[i].value) + " and pdfpath='" + relative_path + "';"
                print(sql)
                cursor.execute(sql)


            #db.commit()
            if round==1:
                sql = "update convert_condition set state=0 where tag=" + str(
                    col_tag[i].value) + " and pdfpath='" + relative_path + "';"
                cursor.execute(sql)
                print(sql)

            db.commit()
            wb.save(in_filename)

        except:
            if round == 1:
                cursor = db.cursor()
                sql = "update convert_condition set state=2 where tag="+str(col_tag[i].value)+" and pdfpath='"+relative_path+"';"
                print(sql)
                cursor.execute(sql)
                db.commit()

    wb.save(in_filename)